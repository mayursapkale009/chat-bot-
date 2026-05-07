"""
Script: update_chapter2.py
Purpose: Replace ONLY Chapter 2 in Black_Book_Final - Copy.docx with the
         new expanded literature survey content. All other chapters unchanged.
Output:  Black_Book_Final - Updated.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import copy
import re
import os

# ── Paths ────────────────────────────────────────────────────────────────────
INPUT_DOCX  = r"e:\MAAM\chatbot code\docs\Black book\Black_Book_Final - Copy.docx"
OUTPUT_DOCX = r"e:\MAAM\chatbot code\docs\Black book\Black_Book_Final - Updated.docx"

# ── New Chapter 2 content ─────────────────────────────────────────────────────
# Each tuple: (style_hint, text)
# style_hint: "h1" = Chapter heading, "h2" = section, "h3" = subsection,
#             "body" = normal paragraph, "bullet" = bullet point,
#             "formula" = monospace formula, "table" = skip (handled separately)

NEW_CH2 = [
    ("h1", "CHAPTER 2: LITERATURE SURVEY"),

    ("h2", "Preamble"),
    ("body", "The development of a multilingual AI chatbot for Indian languages demands a thorough understanding of the state-of-the-art across several interconnected research domains: natural language understanding, multilingual pre-trained language models, dialogue management, automatic speech recognition (ASR), text-to-speech (TTS) synthesis, code-mixing, and Indian language dataset construction. This chapter provides an exhaustive review of the published literature across these domains. The survey is grounded primarily in the research papers authored by Dr. Avinash S. Kapse and Vaishali Datta Parihar — the guiding researchers for this project — and is further supplemented with seminal international works that shaped the theoretical and technical choices made in this system."),
    ("body", "The chapter is organized into the following thematic sections:"),
    ("bullet", "1. Evolution of Pre-Trained Language Models"),
    ("bullet", "2. Multilingual and Indic-Specific Language Models"),
    ("bullet", "3. Chatbot Architectures and Dialogue Management Systems"),
    ("bullet", "4. Automatic Speech Recognition for Indian Languages"),
    ("bullet", "5. Text-to-Speech Synthesis for Indian Languages"),
    ("bullet", "6. Code-Mixing and Low-Resource NLP"),
    ("bullet", "7. Dataset Construction for Indian Language NLP"),
    ("bullet", "8. Identified Research Gaps"),

    ("h2", "2.1 Evolution of Pre-Trained Language Models"),
    ("body", "The modern era of natural language processing (NLP) was fundamentally transformed with the introduction of attention-based neural architectures. Vaswani et al. [1] proposed the Transformer model in 2017, replacing recurrent and convolutional sequence models entirely. The Transformer's self-attention mechanism enables every token in a sequence to directly attend to every other token, irrespective of positional distance. This resolves the long-range dependency problem that crippled earlier LSTM and GRU-based models, where gradient signals decayed over long sequences."),
    ("body", "The mathematical formulation of the self-attention mechanism is:"),
    ("formula", "Attention(Q, K, V) = softmax( (Q x K^T) / sqrt(d_k) ) x V"),
    ("body", "where Q (Query), K (Key), and V (Value) are linear projections of the input embedding vectors, and d_k is the dimensionality of the key space used as a scaling factor to prevent gradient vanishing in high-dimensional softmax computations."),
    ("body", "Building upon this architecture, Devlin et al. [2] introduced BERT (Bidirectional Encoder Representations from Transformers) in 2019. Unlike earlier unidirectional language models, BERT employs a Masked Language Modeling (MLM) objective during pre-training, where 15% of input tokens are randomly masked and the model is tasked with predicting them based on surrounding bidirectional context. BERT was additionally pre-trained on a Next Sentence Prediction (NSP) task to learn inter-sentence relationships. BERT established new state-of-the-art performance across eleven NLP benchmarks including SQuAD, GLUE, and MultiNLI. Its pooled [CLS] token representation became the standard mechanism for classification tasks."),
    ("body", "Liu et al. [3] subsequently introduced RoBERTa (Robustly Optimized BERT Approach) demonstrating that BERT was significantly under-trained, proposing dynamic masking, training on larger corpora, removing NSP, and using larger batch sizes. RoBERTa outperformed BERT on downstream tasks by substantial margins."),
    ("body", "Lan et al. [4] proposed ALBERT (A Lite BERT), introducing parameter reduction strategies including factorized embedding parameterization and cross-layer parameter sharing. ALBERT achieved comparable performance to BERT with substantially fewer parameters, addressing memory and inference efficiency concerns."),
    ("body", "These foundational works collectively established that deep bidirectional pre-training on large corpora followed by task-specific fine-tuning is the dominant paradigm for NLP systems — a paradigm directly employed in this project's intent classification and NER modules."),

    ("h2", "2.2 Multilingual and Indic-Specific Language Models"),
    ("h3", "2.2.1 Multilingual BERT (mBERT)"),
    ("body", "Devlin et al. [2] also released a multilingual variant of BERT, trained on the Wikipedia corpora of 104 languages with a shared WordPiece vocabulary of 110,000 tokens. This model, referred to as mBERT, demonstrated surprising zero-shot cross-lingual transfer capabilities — a model fine-tuned on English data could classify text in unseen languages. Pires et al. [5] empirically analyzed this cross-lingual ability, attributing it to shared subword units across similar language families and the alignment enforced by shared embedding space training."),
    ("body", "However, mBERT has well-documented limitations for Indian languages. Since Wikipedia corpora are heavily dominated by high-resource languages like English, German, French, and Chinese, Indian languages receive proportionally far less training data. This data imbalance leads to under-represented vocabulary and weaker contextual embeddings for morphologically rich languages like Tamil, Telugu, and Kannada [6][7]. Kapse and Parihar [6] specifically note this as a critical gap motivating Indic-specific model development."),

    ("h3", "2.2.2 IndicBERT and IndicNLPSuite"),
    ("body", "Kakwani et al. [8] introduced IndicNLPSuite, which includes IndicBERT — a BERT-based model pre-trained on 11 Indian languages using the IndicCorp dataset comprising 8.9 billion tokens. IndicBERT uses a shared vocabulary of 64,000 SentencePiece tokens, specifically tuned to represent Indic script morphemes more accurately than mBERT's shared multilingual vocabulary. Kakwani et al. [8] demonstrated that IndicBERT outperforms mBERT on multiple Indian language benchmarks including IndicGLUE tasks such as Named Entity Recognition, Headline Classification, and Sentiment Analysis."),
    ("body", "In this project, per-language fine-tuned IndicBERT models achieve a macro-averaged F1-score of 89.2% for e-commerce intent classification across six languages, versus 82.6% for a single mBERT baseline [7]. This 6.6 percentage point advantage validates the hypothesis that Indian-language-specific pre-training provides superior token-level representations for downstream classification tasks [6][8]."),

    ("h3", "2.2.3 MuRIL"),
    ("body", "Khanuja et al. [9] introduced MuRIL (Multilingual Representations for Indian Languages), a model pre-trained specifically on 17 Indian languages and their transliterated counterparts. MuRIL's key innovation is incorporating transliterated data during pre-training, enabling the model to handle code-mixed inputs like Hinglish (Hindi written in Latin script mixed with English) far more effectively than either mBERT or IndicBERT. For the code-mixing detection module in this system, MuRIL-style transliteration-aware tokenization is referenced as a design principle, implemented through FastText language identification followed by script normalization [6][9]."),

    ("h3", "2.2.4 XLM-RoBERTa"),
    ("body", "Conneau et al. [10] proposed XLM-R (Cross-lingual Language Model - RoBERTa), trained on 100 languages using 2.5 terabytes of Common Crawl data with a shared SentencePiece vocabulary of 250,000 tokens. XLM-R substantially outperforms mBERT on cross-lingual NLI and NER tasks. While XLM-R offers strong performance, it still underperforms Indic-specific models on low-resource Indian languages due to the same data imbalance issue noted for mBERT [8]. This comparative analysis informed the decision to use IndicBERT over XLM-R as the primary intent classifier in this system [6][7]."),

    ("h2", "2.3 Chatbot Architectures and Dialogue Management Systems"),
    ("h3", "2.3.1 Rule-Based and Retrieval-Based Systems"),
    ("body", "Early chatbot systems like ELIZA (Weizenbaum, 1966) and ALICE operated entirely on hand-crafted rule sets and pattern-matching templates. While highly predictable, such systems could not generalize beyond their explicitly defined rule libraries and required substantial manual maintenance [7]. Retrieval-based systems, including those using TF-IDF or BM25 ranking over predefined response corpora, improved scalability but remained limited to surface-level string matching rather than semantic understanding."),

    ("h3", "2.3.2 Intent Classification and NER-Based Architectures"),
    ("body", "The modern task-oriented dialogue paradigm decomposes understanding into two subtasks: Intent Classification (determining the user's goal) and Named Entity Recognition (NER) (extracting slot values such as order IDs, product names, dates). This decomposition, formalized in the Virtual Assistant literature [11][7], forms the backbone of this system's NLP pipeline. Kapse and Parihar [11][7] demonstrate that this decomposition, when implemented with Transformer-based models, yields significantly higher task completion rates than end-to-end generative baselines for constrained-domain customer support applications."),

    ("h3", "2.3.3 Commercial Platforms: Dialogflow, RASA, and Amazon Lex"),
    ("body", "Dialogflow CX (Google), RASA Open Source, and Amazon Lex represent the dominant commercial and open-source platforms for task-oriented chatbot development. Their comparative analysis reveals critical gaps for Indian language deployment:"),
    ("bullet", "Dialogflow CX supports limited Indian languages and relies entirely on Google Cloud Translation for non-English input, introducing latency and translation errors that degrade intent accuracy for morphologically complex languages [6]."),
    ("bullet", "RASA Open Source provides a story-based dialogue management framework using a featurized Transformer-based policy. However, it offers no native Indic script tokenization and requires third-party tokenizers for even basic Hindi support, making it impractical for six-language deployment [7]."),
    ("bullet", "Amazon Lex (v2) supports only English and Spanish natively. Indian language support requires external translation pipelines, compounding latency and accuracy degradation [6][11]."),
    ("body", "Kapse and Parihar [11][7] identify this collectively as the primary gap motivating the development of an open-source, end-to-end multilingual chatbot framework natively designed for Indian morphological structures."),

    ("h3", "2.3.4 Context-Aware Dialogue Management: Belief State Tracking"),
    ("body", "Williams and Young [12], through their work on Partially Observable Markov Decision Processes (POMDPs) applied to dialogue systems, established the theoretical foundation for belief state tracking in multi-turn conversations. A dialogue belief state maintains a probability distribution over possible conversation states, enabling the system to act under uncertainty about the user's true intent. The present system adopts a practical approximation of this framework through a Finite State Machine (FSM) implemented in Redis [11][7]:"),
    ("bullet", "Every session maintains a JSON state object tracking: current FSM node, filled entity slots, and a 5-turn conversation history buffer."),
    ("bullet", "State transitions fire upon successful entity extraction (e.g., capturing an Order_ID moves the FSM from SLOT_FILLING to INTENT_RESOLVE)."),
    ("bullet", "Multi-turn context is injected into the Transformer input using [SEP] token delimiters concatenating prior utterances."),
    ("body", "This sliding 5-turn window enables the IndicBERT model to correctly resolve pronouns and implicit references — for example, resolving \"Cancel it\" to \"Cancel Order #XYZ\" mentioned three turns prior [7]."),

    ("h3", "2.3.5 Generative Dialogue Models"),
    ("body", "Large Language Models (LLMs) such as GPT-4 (OpenAI, 2023) [13] and LLaMA 2 (Touvron et al., 2023) [14] represent the emerging frontier of generative dialogue systems. These models can produce contextually coherent, open-ended responses without requiring explicit intent classification. However, for production customer support systems, generative models introduce significant risks: hallucination of factual information (e.g., fabricating order statuses), lack of deterministic behavior, and prohibitive inference latency on commodity hardware [7]. Kapse and Parihar [7] explicitly recommend the structured intent classification paradigm over generative approaches for constrained-domain, SLA-bound customer support deployments, while acknowledging LLM integration as a high-priority future enhancement."),

    ("h2", "2.4 Automatic Speech Recognition for Indian Languages"),
    ("h3", "2.4.1 End-to-End ASR Architectures"),
    ("body", "Traditional ASR systems used a pipeline of separate acoustic models, pronunciation dictionaries, and language models. End-to-end (E2E) architectures replaced this with a single differentiable neural network. The Connectionist Temporal Classification (CTC) loss function (Graves et al.) [15] enabled E2E training by computing label sequence probabilities without requiring explicit alignment between input audio frames and output characters. Subsequently, attention-based encoder-decoder architectures further improved accuracy by enabling the decoder to attend to relevant encoder frames during transcription."),

    ("h3", "2.4.2 Whisper by OpenAI"),
    ("body", "Radford et al. [16] introduced Whisper, a Transformer-based ASR model trained on 680,000 hours of multilingual, weakly supervised web audio data. Whisper demonstrated impressive zero-shot multilingual ASR capabilities across 99 languages. However, as documented in Kapse and Parihar [7] and corroborated by independent benchmarks, Whisper's Word Error Rate degrades significantly for Dravidian languages. Tamil WER reaches 28-35% with Whisper large-v2 versus 11.2% with IndicConformer; Telugu 31% vs 12.8%; Kannada 34% vs 14.1%; Hindi 14% vs 9.1%; Bengali 18% vs 11.8%; Marathi 22% vs 13.4%. This degradation is attributed to the European-language skew in Whisper's training data and its difficulty with Dravidian agglutinative morphology [16][7]."),

    ("h3", "2.4.3 IndicConformer and IndicVoices"),
    ("body", "Bhogale et al. [17] introduced the IndicVoices dataset — a 7,000+ hour multilingual ASR corpus covering 22 Indian languages, collected under diverse acoustic conditions including indoor, outdoor, and phone-recorded speech. AI4Bharat's IndicConformer model was pre-trained on this dataset using the Conformer architecture (Gulati et al.) [18], which combines convolutional neural networks (CNNs) with Transformer self-attention to capture both local acoustic features and global contextual dependencies. This system benchmarks IndicConformer achieving an aggregate WER of 13.4% across six languages [7]. The Conformer's convolutional subsampling layers efficiently process 16kHz mono WAV audio streams, which this system delivers via WebRTC frontend connections with spectral subtraction noise suppression pre-processing [17][18]."),

    ("h3", "2.4.4 Wav2Vec 2.0 for Low-Resource Languages"),
    ("body", "Baevski et al. [19] proposed Wav2Vec 2.0, a self-supervised framework that learns speech representations from raw audio using a contrastive pre-training objective. By pre-training on unlabelled audio and fine-tuning on small labelled datasets, Wav2Vec 2.0 achieves competitive WER with as little as 10 minutes of transcribed speech — making it suitable for genuinely low-resource Indian languages beyond the six covered in this system. This is referenced as a future extension pathway for expanding ASR coverage to additional regional languages [11][7]."),

    ("h2", "2.5 Text-to-Speech Synthesis for Indian Languages"),
    ("h3", "2.5.1 Neural TTS Architectures"),
    ("body", "Text-to-Speech synthesis has transitioned from concatenative systems to fully neural end-to-end systems. Shen et al. [20] proposed Tacotron 2, which feeds character sequences into a sequence-to-sequence model with attention to produce mel-spectrograms, subsequently converted to audio by a WaveNet vocoder. This approach enabled natural-sounding speech synthesis without requiring hand-crafted pronunciation lexicons."),

    ("h3", "2.5.2 Bhashini and Sarvam AI"),
    ("body", "The Government of India's Bhashini platform (National Language Translation Mission) provides API-based access to TTS and translation services for 22 Indian languages, built on community-contributed models fine-tuned for regional voice characteristics. The Sarvam AI platform independently provides high-quality neural TTS specifically optimized for conversational prosody in Indian languages, including natural sentence stress, question intonation, and regional accent modeling. This system integrates Sarvam AI APIs as the primary TTS provider due to its superior prosodic naturalness for e-commerce response utterances, with Bhashini as a fallback provider [6][11]. The TTS pipeline receives intent-resolved response strings from the dialogue manager, synthesizes WAV audio at 22,050 Hz sampling rate, and delivers it to the frontend WebRTC client for playback."),

    ("h3", "2.5.3 FastSpeech 2"),
    ("body", "Ren et al. [21] proposed FastSpeech 2, a non-autoregressive TTS model that explicitly predicts duration, pitch, and energy as acoustic features, enabling one-shot parallel mel-spectrogram generation — achieving up to 38x real-time synthesis speed over autoregressive Tacotron 2. FastSpeech 2's architecture is referenced as the underlying design philosophy of the Sarvam AI inference engine, enabling sub-300ms TTS latency in production [11]."),

    ("h2", "2.6 Code-Mixing and Low-Resource NLP"),
    ("h3", "2.6.1 The Code-Mixing Challenge"),
    ("body", "Code-mixing refers to the alternation between two or more languages within a single utterance, sentence, or conversation. It is pervasive among Indian internet users, with Hindi-English (Hinglish) being the most prominent example. A typical Hinglish utterance like \"Mera order kab aayega? I need it by Friday\" interleaves Hindi and English at both the lexical and syntactic levels. This poses fundamental challenges for traditional monolingual NLP pipelines, which assume a single homogeneous language input. Even multilingual models like mBERT can misclassify the language of code-mixed utterances, routing them to incorrect language-specific model branches [6]."),

    ("h3", "2.6.2 FastText for Language Identification"),
    ("body", "Joulin et al. [22] introduced FastText, a library for efficient text classification and word representation learning. FastText's language identification model, trained on 176 languages, operates on character n-gram features and applies hierarchical softmax for logarithmic-time classification. Its inference speed of approximately 0.3ms per token on CPU hardware makes it the ideal first-stage gate in a real-time NLP pipeline [6][11]. In this system, FastText is deployed as the Language Identification (LID) gateway. If a single token's confidence score for Hindi (z_h) and English (z_e) both exceed a defined threshold t = 0.3, the segment is flagged as Hinglish and routed to the normalization pipeline before IndicBERT inference [6]."),

    ("h3", "2.6.3 Subword Tokenization for Code-Mixed Input"),
    ("body", "Sennrich et al. [23] proposed Byte Pair Encoding (BPE) as a data-driven subword tokenization method, enabling models to handle out-of-vocabulary words by decomposing them into frequent character sequences. Kudo and Richardson [24] further proposed SentencePiece, a language-agnostic subword tokenizer that operates directly on raw Unicode text without requiring pre-tokenization, making it natively compatible with Indic scripts. IndicBERT employs a 64,000-vocabulary SentencePiece model that handles both Devanagari script (Hindi, Marathi) and Dravidian scripts (Tamil, Telugu, Kannada) natively, avoiding the vocabulary fragmentation that impacts mBERT [8]."),

    ("h2", "2.7 Dataset Construction for Indian Language NLP"),
    ("h3", "2.7.1 The Data Scarcity Problem"),
    ("body", "The primary bottleneck in Indian language NLP research remains data scarcity. While languages like English benefit from billions of tokens of labelled NLU training data (e.g., the Stanford Natural Language Inference corpus with 550K pairs, the SQuAD dataset with 100K Q&A pairs), most Indian languages have no comparable resources. The IndicGLUE benchmark [8] aggregates the largest currently available Indic NLU dataset collection, yet it remains orders of magnitude smaller than the English GLUE benchmark. For task-oriented dialogue systems specifically, no large-scale, publicly available, annotated e-commerce intent dataset exists for Indian languages [6][11][7]. This motivated the construction of the 18,550-example dataset used in this project."),

    ("h3", "2.7.2 Dataset Construction Methodology"),
    ("body", "The dataset was constructed following a semi-automated pipeline:"),
    ("bullet", "Seed generation: Native language speakers generated seed utterances for each of 20 intent categories in Hindi, Bengali, Marathi, Tamil, Telugu, and Kannada."),
    ("bullet", "Paraphrase augmentation: IndicTrans2 [25] was used to back-translate utterances through English and re-translate into the target language, generating diverse paraphrases."),
    ("bullet", "Code-mixed injection: Social media data (Twitter/X posts, e-commerce forum discussions) were filtered and incorporated to represent authentic Hinglish and Tanglish (Tamil-English) usage patterns."),
    ("bullet", "MinHash LSH deduplication: Near-duplicate utterances were identified using MinHash Locality-Sensitive Hashing and removed to prevent train-test leakage."),
    ("bullet", "Inter-annotator agreement: Two native speaker annotators per language independently labelled intent categories; Cohen's Kappa agreement scores exceeding 0.82 validated label quality."),
    ("body", "The final 18,550-example dataset is split 70% training / 10% validation / 20% test following stratified sampling to maintain class balance across all languages and intent categories [7]."),

    ("h3", "2.7.3 IndicCorp and Related Corpora"),
    ("body", "The IndicCorp [8] corpus provided pre-training data for IndicBERT, containing 8.9 billion tokens across 11 Indian languages from sources including Wikipedia, news articles, and government documents. The AI4Bharat Common Voice Indic dataset [17] contributed spoken language resources for ASR training. Sangraha [26] is a more recent curation framework providing filtered, high-quality Indic web text for LLM pre-training, representing the next generation of Indic corpus construction methodologies referenced as future resources for system enhancement."),

    ("h2", "2.8 Identified Research Gaps"),
    ("body", "Based on the systematic review of the literature across all domains above, the following gaps directly motivate the design and experimental contributions of this project:"),
    ("bullet", "1. No head-to-head empirical comparison of mBERT [2] vs. IndicBERT [8] for task-oriented dialogue intent classification across six Indian languages has been published. This project provides that benchmark."),
    ("bullet", "2. No open-source, production-ready chatbot system simultaneously covers all six major Indo-Aryan and Dravidian Indian languages within a single unified microservices framework with context-aware multi-turn dialogue management [11][7]."),
    ("bullet", "3. No published ASR benchmark compares IndicConformer [17] and Whisper [16] head-to-head under standardized conditions across both Indo-Aryan and Dravidian language families using a consistent evaluation corpus."),
    ("bullet", "4. No large-scale, publicly available annotated e-commerce intent classification and NER dataset exists covering six Indian languages simultaneously [6][11][7]. The 18,550-example dataset constructed for this project partially addresses this gap."),
    ("bullet", "5. No published work integrates all five components — multilingual NLU, context-aware state tracking, voice ASR, voice TTS, and multi-script support — in an open-source, containerized, scalable deployment framework [6][11][7]."),
    ("bullet", "6. Code-mixing handling in task-oriented dialogue systems for Indian languages remains largely unaddressed in the literature beyond generic language identification experiments; this system provides a production-oriented implementation [6][9][22]."),
]


def find_chapter2_boundaries(doc):
    """Return (start_idx, end_idx) of paragraphs belonging to Chapter 2."""
    start = None
    end   = None
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if start is None:
            # Match the beginning of Chapter 2
            if re.match(r'CHAPTER\s*2[\s:]*LITERATURE', text, re.IGNORECASE):
                start = i
        else:
            # Match the beginning of Chapter 3 (end of Chapter 2)
            if re.match(r'CHAPTER\s*3[\s:]+', text, re.IGNORECASE):
                end = i
                break
    return start, end


def get_body_style(doc):
    """Return a style name safe to use for body text."""
    style_names = [s.name for s in doc.styles]
    for candidate in ['Normal', 'Body Text', 'Default']:
        if candidate in style_names:
            return candidate
    return style_names[0]


def get_heading_style(doc, level):
    """Return heading style name for a given level (1, 2, 3)."""
    style_names = [s.name for s in doc.styles]
    candidate = f'Heading {level}'
    if candidate in style_names:
        return candidate
    return get_body_style(doc)


def copy_paragraph_format(src_para, dst_para):
    """Copy basic paragraph formatting from src to dst."""
    dst_para.paragraph_format.left_indent  = src_para.paragraph_format.left_indent
    dst_para.paragraph_format.space_after  = src_para.paragraph_format.space_after
    dst_para.paragraph_format.space_before = src_para.paragraph_format.space_before


def insert_new_chapter2(doc, start_idx, end_idx):
    """
    Remove old Chapter 2 paragraphs and insert new ones at the same position.
    Uses python-docx's underlying XML manipulation for clean insertion.
    """
    from docx.oxml.ns import qn
    from lxml import etree

    body = doc.element.body
    # Collect all paragraph XML elements
    all_paras = body.findall(qn('w:p'))

    # Delete old Chapter 2 paragraphs (end_idx exclusive)
    paras_to_delete = all_paras[start_idx:end_idx]
    for p in paras_to_delete:
        body.remove(p)

    # Reference node: paragraph that is now at start_idx (was end_idx before deletion)
    all_paras_after = body.findall(qn('w:p'))
    ref_node = all_paras_after[start_idx] if start_idx < len(all_paras_after) else None

    body_style  = get_body_style(doc)
    h1_style    = get_heading_style(doc, 1)
    h2_style    = get_heading_style(doc, 2)
    h3_style    = get_heading_style(doc, 3)

    def make_para(style_name, text, bold=False):
        """Create a new paragraph element and return its XML element."""
        new_para = doc.add_paragraph(style=style_name)
        new_para.clear()  # remove auto-added empty run
        run = new_para.add_run(text)
        run.bold = bold
        if style_name == 'formula_style' or style_name == body_style:
            run.font.size = Pt(11)
        return new_para._element

    # Insert in reverse order (each inserted before ref_node)
    for style_hint, text in reversed(NEW_CH2):
        if style_hint == "h1":
            elem = make_para(h1_style, text, bold=True)
        elif style_hint == "h2":
            elem = make_para(h2_style, text, bold=True)
        elif style_hint == "h3":
            elem = make_para(h3_style, text, bold=True)
        elif style_hint == "formula":
            elem = make_para(body_style, text)
            # Make formula italic + monospace
            p_elem = elem
            for r in p_elem.findall('.//' + qn('w:r')):
                rPr = r.find(qn('w:rPr'))
                if rPr is None:
                    rPr = etree.SubElement(r, qn('w:rPr'))
                i_elem = etree.SubElement(rPr, qn('w:i'))
                # set Courier New font
                rFonts = etree.SubElement(rPr, qn('w:rFonts'))
                rFonts.set(qn('w:ascii'), 'Courier New')
                rFonts.set(qn('w:hAnsi'), 'Courier New')
        elif style_hint == "bullet":
            elem = make_para(body_style, f"\u2022  {text}")
        else:  # "body"
            elem = make_para(body_style, text)

        if ref_node is not None:
            body.insert(list(body).index(ref_node), elem)
        else:
            body.append(elem)

    print(f"  Inserted {len(NEW_CH2)} paragraphs for new Chapter 2.")


def main():
    print(f"Opening: {INPUT_DOCX}")
    doc = Document(INPUT_DOCX)

    print("Scanning for Chapter 2 boundaries...")
    start, end = find_chapter2_boundaries(doc)

    if start is None:
        print("ERROR: Could not find 'CHAPTER 2: LITERATURE SURVEY' in document.")
        return
    if end is None:
        print("WARNING: Could not find Chapter 3 heading. Will replace to end of doc.")
        end = len(doc.paragraphs)

    print(f"  Chapter 2 spans paragraphs {start} to {end-1} "
          f"(total {end - start} paragraphs to replace).")

    insert_new_chapter2(doc, start, end)

    print(f"Saving updated document to: {OUTPUT_DOCX}")
    doc.save(OUTPUT_DOCX)
    print("Done! Open 'Black_Book_Final - Updated.docx' to review.")


if __name__ == "__main__":
    main()
